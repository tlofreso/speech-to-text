from openai import OpenAI
import dropbox
import os, sys, json
from rich import print
from pydub import AudioSegment
from pydub.utils import make_chunks
from meeting_notes import meeting_minutes
from docx import Document

def get_config():
    required_vars = ["DROPBOX_APP_KEY", "DROPBOX_APP_SECRET", "DROPBOX_REFRESH_TOKEN", "OPENAI_API_KEY"]
    config = {}

    missing_vars = [var for var in required_vars if var not in os.environ]
    if missing_vars:
        sys.exit(f"Error: Missing environment variables: {', '.join(missing_vars)}")

    for var in required_vars:
        config[var.lower()] = os.environ[var]

    return config

config = get_config()
openai_client = OpenAI()
dropbox_client = dropbox.Dropbox(app_key=
                                 config['dropbox_app_key'],
                                 app_secret=config['dropbox_app_secret'],
                                 oauth2_refresh_token=config['dropbox_refresh_token']
)
AUDIO_PATH = "voice-memos"
TEXT_PATH = "text-transcripts"

def get_filenames(path):
    filenames = []
    print("Getting files...")
    files = dropbox_client.files_list_folder(path=f"/{path}")
    if len(files.entries) < 1:
        print("no files found")
        sys.exit()
    for file in files.entries:
        filenames.append(file.name)

    print(f"Found {len(files.entries)} files...")
    return filenames

def download_files(path, files):
    for file in files:
        print(f"Downloading audio files...")
        dropbox_client.files_download_to_file(file, f"/{path}/{file}")
    print("Done.")

def upload_file(path, file):
    print(f"Uploading text file...")
    with open(file, "rb") as f:
        data = f.read()
    dropbox_client.files_upload(data, path=f"/{path}/{file}")

def delete_file(path, file):
    print(f"Deleting audio file...")
    dropbox_client.files_delete(f"/{path}/{file}")

def transcribe_audio(file, prompt=None):
    print("Starting transcribe...")
    with open(file, "rb") as file:
        transcript = openai_client.audio.transcriptions.create(model = "whisper-1", file=file, prompt=prompt)
    
    return transcript.text

def chunk_audio(chunk_mins, file):
    print("Chunking audio...")

    chunks_out = []

    audio = AudioSegment.from_file(file)
    chunk_size = chunk_mins * 60 * 1000
    chunks = make_chunks(audio, chunk_size)
    for i, chunk in enumerate(chunks):
        audio_out = f'{i}.mp3'
        chunk.export(audio_out, format='mp3')
        chunks_out.append(audio_out)
    
    print(f"Chunked into {len(chunks_out)} files...")
    return(chunks_out)


def save_as_docx(minutes, filename):
    doc = Document()
    for key, value in minutes.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    doc.save(filename)

if __name__ == "__main__":

    audio_files = get_filenames(AUDIO_PATH)
    download_files(AUDIO_PATH, audio_files)

    for file in audio_files:

        chunks = chunk_audio(20, file)
        transcripts = ['']

        for chunk in chunks:
            transcript = transcribe_audio(chunk, transcripts[-1])
            transcripts.append(transcript)

        out_txt = file.split('.')[0] + ".txt"
        with open(f"{out_txt}", "w") as output_file:
            output_file.write(' '.join(transcripts))
        
        print(f"Wrote {out_txt}...")
        
        out_docx = file.split('.')[0] + ".docx"
        with open(out_txt, "r") as full_transcript:
            notes = meeting_minutes(full_transcript.read())
        
        save_as_docx(notes, out_docx)
        
        upload_file(f"{TEXT_PATH}/{file.split('.')[0]}", out_txt)
        upload_file(f"{TEXT_PATH}/{file.split('.')[0]}", out_docx)
        delete_file(AUDIO_PATH, file)
    
    print("Done.")

    

